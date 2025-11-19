import openpyxl
from docx import Document
import os

# === RUTAS ===
ruta_excel = r"C:\Users\CamiloAmaya\OneDrive - Glenfarne Companies\EXCEL_TEST.xlsx"
ruta_word = r"C:\Users\CamiloAmaya\OneDrive - Glenfarne Companies\TF0.docx"
ruta_salida = r"C:\Users\CamiloAmaya\OneDrive - Glenfarne Companies\Informes_generados"

# Crear carpeta de salida si no existe
os.makedirs(ruta_salida, exist_ok=True)

# === CARGAR EXCEL ===
wb = openpyxl.load_workbook(ruta_excel)
hoja = wb.active  # Usa la hoja activa o cámbiala por el nombre exacto, ej: wb["Hoja1"]

# === RECORRER FILAS ===
for fila in hoja.iter_rows(min_row=2, values_only=True):  # desde la fila 2 si hay encabezados
    task = fila[6]  # Columna G (índice 6 porque empieza en 0)

    # Saltar filas vacías
    if not task:
        continue

    # === CARGAR PLANTILLA WORD ===
    doc = Document(ruta_word)

    # Reemplazar {{task}} en párrafos
    for p in doc.paragraphs:
        if "{{task}}" in p.text:
            p.text = p.text.replace("{{task}}", str(task))

    # Reemplazar {{task}} dentro de tablas
    for tabla in doc.tables:
        for fila_t in tabla.rows:
            for celda in fila_t.cells:
                if "{{task}}" in celda.text:
                    celda.text = celda.text.replace("{{task}}", str(task))

    # === GUARDAR NUEVO DOCUMENTO ===
    nombre_archivo = f"TASK_{task}_TEST.docx"
    ruta_guardado = os.path.join(ruta_salida, nombre_archivo)
    doc.save(ruta_guardado)

    print(f"✅ Documento generado: {ruta_guardado}")

print("🎉 Prueba completada. Se generaron los archivos correctamente.")
